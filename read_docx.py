import docx
doc = docx.Document(r'C:\Users\MR_V\AppData\Roaming\Minetest\sora_story\main story_elementary_1.docx')
with open('story_output.txt', 'w', encoding='utf-8') as f:
    f.write(f'Paragraphs: {len(doc.paragraphs)}\n')
    for i, para in enumerate(doc.paragraphs):
        f.write(f'{i}: "{para.text}"\n')
    f.write('---\n')
    f.write(f'Tables: {len(doc.tables)}\n')
    for t_idx, table in enumerate(doc.tables):
        f.write(f'Table {t_idx}:\n')
        for row in table.rows:
            f.write(str([cell.text for cell in row.cells]) + '\n')